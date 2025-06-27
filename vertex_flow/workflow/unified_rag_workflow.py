#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
统一配置的RAG Workflow实现
基于vertex flow接口和统一配置实现的本地检索增强生成系统
"""

import os
import time
from pathlib import Path
from typing import Any, Dict, List, Optional

from ..utils.logger import LoggerUtil
from .constants import SYSTEM, USER
from .service import VertexFlowService
from .vertex import (
    EmbeddingVertex,
    FunctionVertex,
    LLMVertex,
    SinkVertex,
    SourceVertex,
    VectorQueryVertex,
    VectorStoreVertex,
)
from .vertex.embedding_providers import TextEmbeddingProvider
from .vertex.vector_engines import Doc, VectorEngine
from .vertex.vertex_group import VertexGroup
from .workflow import Workflow, WorkflowContext

logging = LoggerUtil.get_logger()


class ResultFormatterVertex(FunctionVertex):
    """结果格式化顶点，用于格式化向量检索结果"""

    def __init__(
        self,
        id: str,
        name: Optional[str] = None,
        params: Optional[Dict[str, Any]] = None,
        variables: Optional[List[Dict[str, Any]]] = None,
    ):
        super().__init__(
            id=id,
            name=name or "结果格式化",
            task=self.format_results,
            params=params or {},
            variables=variables or [],
        )

    def format_results(self, inputs: Dict[str, Any], context=None):
        """
        格式化向量检索结果

        Args:
            inputs: 包含检索结果的输入

        Returns:
            格式化后的结果
        """
        results = inputs.get("results", [])
        logging.info(f"results: {results}")

        if not results:
            return {"formatted_results": "未找到相关文档。"}

        formatted_results = []
        for i, result in enumerate(results, 1):
            content = result.get("content", "")
            score = result.get("score", 0)
            doc_id = result.get("id", f"文档{i}")

            # 格式化每个检索结果
            formatted_result = f"文档{i} (ID: {doc_id}, 相似度: {score:.3f}):\n{content}\n"
            formatted_results.append(formatted_result)

        formatted_text = "\n".join(formatted_results)
        return {"formatted_results": formatted_text}


class EnhancedResultFormatterVertex(FunctionVertex):
    """增强的结果格式化顶点，支持Web搜索结果的URL和summary信息"""

    def __init__(
        self,
        id: str,
        name: Optional[str] = None,
        params: Optional[Dict[str, Any]] = None,
        variables: Optional[List[Dict[str, Any]]] = None,
    ):
        super().__init__(
            id=id,
            name=name or "增强结果格式化",
            task=self.format_results_with_web_search,
            params=params or {},
            variables=variables or [],
        )

    def format_results_with_web_search(self, inputs: Dict[str, Any], context=None):
        """
        格式化向量检索结果，并整合Web搜索结果的URL和summary信息

        Args:
            inputs: 包含检索结果和可能的Web搜索结果的输入

        Returns:
            格式化后的结果，包含参考链接和总结信息
        """
        results = inputs.get("results", [])
        web_search_results = inputs.get("web_search_results", [])
        web_search_summary = inputs.get("web_search_summary", "")
        
        logging.info(f"Vector results: {len(results)}, Web search results: {len(web_search_results)}")

        formatted_sections = []
        
        # 格式化向量检索结果
        if results:
            vector_results = []
            for i, result in enumerate(results, 1):
                content = result.get("content", "")
                score = result.get("score", 0)
                doc_id = result.get("id", f"文档{i}")
                
                formatted_result = f"文档{i} (ID: {doc_id}, 相似度: {score:.3f}):\n{content}\n"
                vector_results.append(formatted_result)
            
            formatted_sections.append("## 本地知识库检索结果\n" + "\n".join(vector_results))
        
        # 格式化Web搜索结果
        if web_search_results:
            web_results = []
            reference_links = []
            
            for i, result in enumerate(web_search_results, 1):
                title = result.get("title", f"网络结果{i}")
                url = result.get("url", "")
                snippet = result.get("snippet", "")
                source = result.get("source", "Web")
                
                # 格式化Web搜索结果
                web_result = f"网络结果{i} ({source}):\n标题: {title}\n内容: {snippet}\n"
                if url:
                    web_result += f"链接: {url}\n"
                    reference_links.append(f"[{i}] {title}: {url}")
                
                web_results.append(web_result)
            
            formatted_sections.append("## 网络搜索结果\n" + "\n".join(web_results))
            
            # 添加参考链接部分
            if reference_links:
                formatted_sections.append("## 参考链接\n" + "\n".join(reference_links))
        
        # 添加Web搜索总结
        if web_search_summary:
            formatted_sections.append(f"## 网络搜索总结\n{web_search_summary}")
        
        # 如果没有任何结果
        if not formatted_sections:
            return {"formatted_results": "未找到相关文档或网络信息。"}
        
        formatted_text = "\n\n".join(formatted_sections)
        return {"formatted_results": formatted_text}


class WebSearchIntegrationVertex(FunctionVertex):
    """Web搜索集成顶点，用于在RAG流程中执行Web搜索"""

    def __init__(
        self,
        id: str,
        name: Optional[str] = None,
        params: Optional[Dict[str, Any]] = None,
        variables: Optional[List[Dict[str, Any]]] = None,
        web_search_service=None,
    ):
        super().__init__(
            id=id,
            name=name or "Web搜索集成",
            task=self.perform_web_search,
            params=params or {},
            variables=variables or [],
        )
        self.web_search_service = web_search_service
        self.search_count = params.get("search_count", 3) if params else 3
        self.enable_summary = params.get("enable_summary", True) if params else True

    def perform_web_search(self, inputs: Dict[str, Any], context=None):
        """
        执行Web搜索并提取结果

        Args:
            inputs: 包含搜索查询的输入

        Returns:
            Web搜索结果，包含URL和summary信息
        """
        # 获取搜索查询
        search_queries = inputs.get("search_queries", [])
        original_query = inputs.get("original_query", "")
        
        # 如果没有专门的搜索查询，使用原始查询
        if not search_queries and original_query:
            search_queries = [original_query]
        
        if not search_queries:
            logging.warning("没有提供搜索查询，跳过Web搜索")
            return {
                "web_search_results": [],
                "web_search_summary": "",
                "search_performed": False
            }
        
        # 如果没有Web搜索服务，跳过搜索
        if not self.web_search_service:
            logging.info("Web搜索服务未配置，跳过网络搜索")
            return {
                "web_search_results": [],
                "web_search_summary": "",
                "search_performed": False
            }
        
        all_results = []
        all_summaries = []
        
        # 对每个搜索查询执行搜索
        for query in search_queries[:2]:  # 限制最多2个查询以避免过多请求
            try:
                logging.info(f"执行Web搜索: {query}")
                
                # 调用Web搜索工具
                search_inputs = {
                    "query": query,
                    "count": self.search_count,
                    "summary": self.enable_summary
                }
                
                search_result = self.web_search_service.func(search_inputs)
                
                if search_result.get("success", False):
                    results = search_result.get("results", [])
                    summary = search_result.get("summary", "")
                    
                    all_results.extend(results)
                    if summary:
                        all_summaries.append(f"查询'{query}': {summary}")
                    
                    logging.info(f"Web搜索成功，查询: {query}, 结果数: {len(results)}")
                else:
                    error = search_result.get("error", "未知错误")
                    logging.warning(f"Web搜索失败，查询: {query}, 错误: {error}")
                    
            except Exception as e:
                logging.error(f"Web搜索异常，查询: {query}, 错误: {str(e)}")
                continue
        
        # 去重结果（基于URL）
        seen_urls = set()
        unique_results = []
        for result in all_results:
            url = result.get("url", "")
            if url and url not in seen_urls:
                seen_urls.add(url)
                unique_results.append(result)
            elif not url:  # 没有URL的结果也保留
                unique_results.append(result)
        
        # 合并总结
        combined_summary = "\n".join(all_summaries) if all_summaries else ""
        
        logging.info(f"Web搜索完成，总结果数: {len(unique_results)}, 有总结: {bool(combined_summary)}")
        
        return {
            "web_search_results": unique_results,
            "web_search_summary": combined_summary,
            "search_performed": len(unique_results) > 0
        }


class DocumentProcessorVertex(FunctionVertex):
    """文档处理顶点"""

    def __init__(
        self,
        id: str,
        name: Optional[str] = None,
        params: Optional[Dict[str, Any]] = None,
        variables: Optional[List[Dict[str, Any]]] = None,
        vector_engine=None,
    ):
        super().__init__(
            id=id,
            name=name or "文档处理",
            task=self.process_documents,
            params=params or {},
            variables=variables or [],
        )
        self.chunk_size = params.get("chunk_size", 1000) if params else 1000
        self.chunk_overlap = params.get("chunk_overlap", 200) if params else 200
        self.vector_engine = vector_engine

    def _load_existing_hashes(self) -> Dict[str, Dict]:
        """
        加载已处理文件的哈希内容

        Returns:
            包含内容哈希和元数据的字典
        """
        if self.vector_engine and hasattr(self.vector_engine, "content_hashes"):
            return self.vector_engine.content_hashes
        return {}

    def _get_document_hashes(self, docs: List) -> Dict[str, Dict]:
        """
        获取当前要处理文档的哈希内容

        Args:
            docs: 文档列表

        Returns:
            文档路径/ID到哈希和内容的映射
        """
        doc_hashes = {}

        for i, doc in enumerate(docs):
            doc_key = None
            content = None
            metadata = {}

            if isinstance(doc, str):
                if os.path.exists(doc):
                    # 文件路径
                    doc_key = doc
                    content = self._load_file(doc)
                    if content and os.path.exists(doc):
                        metadata = {"source": doc, "mtime": os.path.getmtime(doc), "file_size": os.path.getsize(doc)}
                else:
                    # 文本内容
                    doc_key = f"text_{i}"
                    content = doc
                    metadata = {"source": "text", "mtime": 0}

            elif isinstance(doc, dict):
                content = doc.get("content", "")
                doc_key = doc.get("metadata", {}).get("source", f"dict_{i}")
                metadata = doc.get("metadata", {})

            if content and doc_key:
                source_path = metadata.get("source")
                file_path = source_path if source_path and source_path != "text" else None
                content_hash = self._generate_content_hash(content, file_path=file_path)
                doc_hashes[doc_key] = {
                    "hash": content_hash,
                    "content": content,
                    "metadata": metadata,
                    "original_index": i,
                }

        return doc_hashes

    def _filter_duplicates(self, doc_hashes: Dict[str, Dict], existing_hashes: Dict[str, Dict]) -> tuple:
        """
        过滤重复文档

        Args:
            doc_hashes: 当前文档哈希
            existing_hashes: 已存在的哈希

        Returns:
            (需要处理的文档, 跳过的文档列表)
        """
        to_process = {}
        skipped = []

        for doc_key, doc_info in doc_hashes.items():
            content_hash = doc_info["hash"]

            if content_hash in existing_hashes:
                # 内容重复，跳过
                skipped.append(doc_key)
            else:
                # 新内容，需要处理
                to_process[doc_key] = doc_info

        return to_process, skipped

    def _process_unique_documents(self, to_process: Dict[str, Dict]) -> List[Dict]:
        """
        处理非重复文档，进行分块

        Args:
            to_process: 需要处理的文档

        Returns:
            处理后的文档块列表
        """
        processed_docs = []

        for doc_key, doc_info in to_process.items():
            content = doc_info["content"]
            metadata = doc_info["metadata"]
            original_index = doc_info["original_index"]

            # 对文档进行分块
            chunks = self._chunk_text(content)

            for j, chunk in enumerate(chunks):
                processed_docs.append(
                    {
                        "id": f"doc_{original_index}_chunk_{j}",
                        "content": chunk,
                        "metadata": {
                            **metadata,
                            "chunk_index": j,
                            "total_chunks": len(chunks),
                            "length": len(chunk),
                        },
                    }
                )

        return processed_docs

    def _update_vector_hashes(self, doc_hashes: Dict[str, Dict]):
        """
        将新的文档哈希信息更新到向量引擎中

        Args:
            doc_hashes: 文档哈希信息
        """
        if self.vector_engine and hasattr(self.vector_engine, "content_hashes"):
            for doc_key, doc_info in doc_hashes.items():
                content_hash = doc_info["hash"]
                metadata = doc_info["metadata"]
                self.vector_engine.content_hashes[content_hash] = {
                    "id": f"doc_{doc_info['original_index']}",
                    "metadata": metadata,
                }

    def process_documents(self, inputs: Dict[str, Any], context=None):
        """
        处理文档，支持多种格式和分块，并在构建索引前检查重复

        Args:
            inputs: 包含文档路径或内容的输入

        Returns:
            处理后的文档列表
        """
        local_inputs = self.resolve_dependencies(inputs=inputs)
        docs = local_inputs.get("docs", [])

        if isinstance(docs, str):
            docs = [docs]

        # 1. 先加载以往处理过的文件的哈希内容
        existing_hashes = self._load_existing_hashes()

        # 2. 对这次要处理的文件获取哈希内容
        doc_hashes = self._get_document_hashes(docs)

        # 3. 如果有相同的，则跳过；如果不同，继续处理
        to_process, skipped = self._filter_duplicates(doc_hashes, existing_hashes)

        # 4. 处理非重复文档
        processed_docs = self._process_unique_documents(to_process)

        # 5. 更新向量引擎中的哈希信息，避免以后再重复处理
        self._update_vector_hashes(to_process)

        # 日志输出
        if len(skipped) > 0:
            logging.info(f"文档处理完成: 生成 {len(processed_docs)} 个块，跳过 {len(skipped)} 个重复文档")
        else:
            logging.info(f"文档处理完成: 生成 {len(processed_docs)} 个块")

        # 如果所有文档都被去重了，返回空列表表示跳过后续处理
        if not processed_docs:
            logging.info("所有文档都已存在，跳过embedding和向量存储步骤")
            return {"docs": []}

        return {"docs": processed_docs}

    def _load_file(self, file_path: str) -> Optional[str]:
        """加载文件内容，支持多种编码兜底"""
        try:
            file_ext = Path(file_path).suffix.lower()

            if file_ext in {".txt", ".md"}:
                # 尝试多种编码方式
                for encoding in ["utf-8", "gbk", "gb2312", "big5", "utf-16", "latin-1"]:
                    try:
                        with open(file_path, "r", encoding=encoding) as f:
                            content = f.read()
                            logging.info(f"成功使用 {encoding} 编码读取文件: {file_path}")
                            return content
                    except UnicodeDecodeError:
                        continue
                    except Exception:
                        break

                # 如果所有编码都失败，尝试二进制读取并强制解码
                try:
                    with open(file_path, "rb") as f:
                        raw_content = f.read()
                        content = raw_content.decode("utf-8", errors="ignore")
                        logging.warning(f"使用忽略错误的 UTF-8 解码读取文件: {file_path}")
                        return content
                except Exception:
                    logging.error(f"无法以任何编码方式读取文件: {file_path}")
                    return None

            elif file_ext == ".pdf":
                return self._load_pdf(file_path)

            elif file_ext in {".docx", ".doc"}:
                return self._load_word(file_path)

            else:
                logging.warning(f"不支持的文件格式: {file_ext}")
                return None

        except Exception as e:
            logging.error(f"加载文件 {file_path} 失败: {e}")
            return None

    def _load_pdf(self, file_path: str) -> Optional[str]:
        """加载PDF文件"""
        try:
            import PyPDF2

            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                # 清理PDF文本，去除多余的空白字符
                text = self._clean_pdf_text(text)
                return text
        except ImportError:
            logging.error("请安装PyPDF2: pip install PyPDF2")
            return None
        except Exception as e:
            logging.error(f"PDF加载失败: {e}")
            return None

    def _clean_pdf_text(self, text: str) -> str:
        """清理PDF文本，去除多余的空白字符和换行符"""
        import re

        # 去除多余的空白字符
        text = re.sub(r"\s+", " ", text)
        # 去除行首行尾空白
        text = text.strip()
        # 将多个换行符替换为单个换行符
        text = re.sub(r"\n+", "\n", text)

        return text

    def _load_word(self, file_path: str) -> Optional[str]:
        """加载Word文件"""
        try:
            from docx import Document

            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            logging.error("请安装python-docx: pip install python-docx")
            return None
        except Exception as e:
            logging.error(f"Word文档加载失败: {e}")
            return None

    def _chunk_text(self, text: str) -> List[str]:
        """
        将文本分块

        Args:
            text: 输入文本

        Returns:
            分块后的文本列表
        """
        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size

            # 如果不是最后一块，尝试在句子边界分割
            if end < len(text):
                # 寻找最近的句子结束符
                sentence_end = max(
                    text.rfind(".", start, end),
                    text.rfind("!", start, end),
                    text.rfind("?", start, end),
                    text.rfind("\n", start, end),
                )

                if sentence_end > start:
                    end = sentence_end + 1

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # 计算下一个块的起始位置，考虑重叠
            start = max(start + 1, end - self.chunk_overlap)

        return chunks

    def _generate_content_hash(self, content: str, file_path: Optional[str] = None) -> str:
        """生成文档内容的哈希值，pdf/doc/docx用二进制hash，其它用文本hash"""
        import hashlib
        import os
        from pathlib import Path

        # 如果是PDF或Word等二进制文件，直接用文件二进制内容hash
        if file_path and Path(file_path).suffix.lower() in {".pdf", ".docx", ".doc"} and os.path.exists(file_path):
            with open(file_path, "rb") as f:
                file_bytes = f.read()
            return hashlib.md5(file_bytes).hexdigest()
        # 否则用文本内容hash
        cleaned_content = self._clean_pdf_text(content).lower()
        return hashlib.md5(cleaned_content.encode("utf-8")).hexdigest()

    def _is_duplicate_content(self, content_hash: str) -> bool:
        """检查文档内容是否已存在（需要从向量引擎获取）"""
        # 这里需要从向量引擎获取内容哈希信息
        # 由于DocumentProcessorVertex无法直接访问向量引擎，我们暂时返回False
        # 实际的去重逻辑在VectorStoreVertex中处理
        return False

    def _get_existing_metadata(self, content_hash: str) -> Optional[Dict]:
        """获取已存在文档的元数据（需要从向量引擎获取）"""
        # 这里需要从向量引擎获取元数据信息
        # 由于DocumentProcessorVertex无法直接访问向量引擎，我们暂时返回None
        # 实际的元数据获取逻辑在VectorStoreVertex中处理
        return None


class UnifiedRAGWorkflowBuilder:
    """统一配置的RAG工作流构建器"""

    def __init__(self):
        """
        初始化统一配置的RAG工作流构建器
        """
        # 使用正确的配置文件路径
        # config_path = os.path.join(os.getcwd(), "config", "llm.yml")
        # 复用VertexFlowService的配置
        self.service = VertexFlowService()
        self.config: Dict[str, Any] = self.service._config
        self.embedding_provider = None
        self.vector_engine = None
        self.llm_model = None
        self.web_search_service = None
        self._initialize_components()

    def _format_search_results(self, results):
        """格式化搜索结果，提取有用的信息"""
        if not results:
            return "未找到相关文档。"

        formatted_results = []
        for i, result in enumerate(results, 1):
            content = result.get("content", "")
            score = result.get("score", 0)
            doc_id = result.get("id", f"文档{i}")

            # 格式化每个检索结果
            formatted_result = f"文档{i} (ID: {doc_id}, 相似度: {score:.3f}):\n{content}\n"
            formatted_results.append(formatted_result)

        return "\n".join(formatted_results)

    def _get_default_config(self) -> Dict[str, Any]:
        """获取默认配置"""
        return {
            "embedding": {
                "local": {"enabled": True, "model_name": "all-MiniLM-L6-v2", "dimension": 384},
                "dashscope": {"enabled": False},
                "bce": {"enabled": False},
            },
            "vector": {
                "local": {"enabled": True, "dimension": 384, "index_name": "default"},
                "dashvector": {"enabled": False},
            },
            "llm": {"deepseek": {"enabled": True}},
            "document": {"chunk_size": 1000, "chunk_overlap": 200},
            "retrieval": {"top_k": 3, "similarity_threshold": 0.3, "rerank": False},
        }

    def _initialize_components(self):
        """初始化各个组件"""
        # 初始化服务
        self.service = VertexFlowService()

        # 初始化向量引擎
        try:
            self.vector_engine = self.service.get_smart_vector_engine()
            logging.info(f"使用向量引擎: {self.vector_engine.__class__.__name__}")
        except Exception as e:
            logging.error(f"无法初始化向量引擎: {e}")
            raise ValueError("无法初始化向量引擎，请检查配置")

        # 初始化嵌入提供者
        try:
            self.embedding_provider = self.service.get_embedding()
            logging.info(f"使用嵌入提供者: {self.embedding_provider.__class__.__name__}")
        except Exception as e:
            logging.error(f"无法初始化嵌入提供者: {e}")
            raise ValueError("无法初始化嵌入提供者，请检查配置")

        # 初始化LLM模型
        try:
            self.llm_model = self.service.get_chatmodel()
            logging.info(f"使用LLM模型: {self.llm_model.__class__.__name__}")
        except Exception as e:
            logging.error(f"无法初始化LLM模型: {e}")
            raise ValueError("无法初始化LLM模型，请检查配置")

        # 初始化网络搜索服务
        try:
            self.web_search_service = self.service.get_web_search_tool()
            logging.info(f"使用网络搜索服务: {self.web_search_service.name}")
        except Exception as e:
            logging.warning(f"无法初始化网络搜索服务: {e}")
            logging.info("查询改写功能将仅使用LLM，不包含网络搜索")
            self.web_search_service = None

    def build_smart_indexing_workflow(self) -> Workflow:
        """构建智能索引工作流，在构建索引前先检查重复"""
        document_config = self.config.get("document", {})

        # 创建工作流上下文
        context = WorkflowContext()

        # 创建顶点
        source_vertex = SourceVertex(id="SOURCE", name="文档输入", task=lambda inputs, context: inputs)

        # 智能文档处理器，在构建索引前检查重复
        smart_doc_processor = DocumentProcessorVertex(
            id="SMART_DOC_PROCESSOR",
            name="智能文档处理",
            params=document_config if isinstance(document_config, dict) else {},
            variables=[
                {
                    "source_scope": "SOURCE",
                    "source_var": "docs",
                    "local_var": "docs",
                }
            ],
            vector_engine=self.vector_engine,
        )

        embedding_vertex = EmbeddingVertex(
            id="EMBEDDING",
            name="文档向量化",
            embedding_provider=self.embedding_provider,
            variables=(
                [
                    {
                        "source_scope": "SMART_DOC_PROCESSOR",
                        "source_var": "docs",
                        "local_var": "docs",
                    }
                ]
                if self.embedding_provider
                else None
            ),
        )

        vector_store = VectorStoreVertex(
            id="VECTOR_STORE",
            name="向量存储",
            vector_engine=self.vector_engine,
            variables=[
                {
                    "source_scope": "EMBEDDING",
                    "source_var": "embeddings",
                    "local_var": "docs",
                }
            ],
        )

        sink_vertex = SinkVertex(id="SINK", name="完成", task=lambda inputs, context: logging.info("智能索引构建完成"))

        # 创建工作流
        workflow = Workflow(context=context)
        workflow.add_vertex(source_vertex)
        workflow.add_vertex(smart_doc_processor)
        workflow.add_vertex(embedding_vertex)
        workflow.add_vertex(vector_store)
        workflow.add_vertex(sink_vertex)

        # 构建工作流图
        source_vertex | smart_doc_processor | embedding_vertex | vector_store | sink_vertex

        return workflow

    def _smart_process_documents(self, inputs: Dict[str, Any], context=None):
        """
        智能文档处理，只根据内容hash去重，hash已存在则跳过embedding
        """
        docs = inputs.get("docs", [])
        if isinstance(docs, str):
            docs = [docs]
        processed_docs = []
        skipped_docs = []
        content_hashes = {}
        if hasattr(self.vector_engine, "content_hashes"):
            content_hashes = self.vector_engine.content_hashes

        # 记录开始处理
        logging.info(f"开始智能处理 {len(docs)} 个文档")

        for i, doc in enumerate(docs):
            if isinstance(doc, str):
                if os.path.exists(doc):
                    content = self._load_file_content(doc)
                    if content:
                        content_hash = self._generate_content_hash(content, file_path=doc)
                        if content_hash in content_hashes:
                            skipped_docs.append(doc)
                            continue
                        processed_docs.extend(self._process_single_document(doc, content, i))
                else:
                    content = doc
                    if content:
                        content_hash = self._generate_content_hash(content)
                        if content_hash in content_hashes:
                            skipped_docs.append(f"text_{i}")
                            continue
                        processed_docs.extend(self._process_single_document(f"text_{i}", content, i))
            elif isinstance(doc, dict):
                content = doc.get("content", "")
                file_path = doc.get("metadata", {}).get("source", None)
                if content:
                    content_hash = self._generate_content_hash(content, file_path=file_path)
                    if content_hash in content_hashes:
                        skipped_docs.append(doc.get("metadata", {}).get("source", f"dict_{i}"))
                        continue
                    processed_docs.append(doc)
                else:
                    processed_docs.append(doc)

        # 简化日志输出，只显示关键统计信息
        if len(skipped_docs) > 0:
            logging.info(f"文档处理完成: 生成 {len(processed_docs)} 个块，跳过 {len(skipped_docs)} 个重复文档")
        else:
            logging.info(f"文档处理完成: 生成 {len(processed_docs)} 个块")

        return {"docs": processed_docs}

    def _process_single_document(self, doc_path: str, content: str, doc_index: int, force_update: bool = False):
        """处理单个文档，返回文档块列表"""
        chunks = self._chunk_text(content)
        processed_chunks = []

        for j, chunk in enumerate(chunks):
            metadata = {
                "source": doc_path,
                "chunk_index": j,
                "total_chunks": len(chunks),
                "length": len(chunk),
            }

            if os.path.exists(doc_path):
                metadata["mtime"] = os.path.getmtime(doc_path)
                metadata["file_size"] = os.path.getsize(doc_path)

            if force_update:
                metadata["force_update"] = True

            processed_chunks.append(
                {
                    "id": f"doc_{doc_index}_chunk_{j}",
                    "content": chunk,
                    "metadata": metadata,
                }
            )

        return processed_chunks

    def _load_file_content(self, file_path: str) -> Optional[str]:
        """加载文件内容，支持多种编码兜底"""
        try:
            file_ext = Path(file_path).suffix.lower()

            if file_ext in {".txt", ".md"}:
                # 尝试多种编码方式
                for encoding in ["utf-8", "gbk", "gb2312", "big5", "utf-16", "latin-1"]:
                    try:
                        with open(file_path, "r", encoding=encoding) as f:
                            content = f.read()
                            logging.info(f"成功使用 {encoding} 编码读取文件: {file_path}")
                            return content
                    except UnicodeDecodeError:
                        continue
                    except Exception:
                        break

                # 如果所有编码都失败，尝试二进制读取并强制解码
                try:
                    with open(file_path, "rb") as f:
                        raw_content = f.read()
                        content = raw_content.decode("utf-8", errors="ignore")
                        logging.warning(f"使用忽略错误的 UTF-8 解码读取文件: {file_path}")
                        return content
                except Exception:
                    logging.error(f"无法以任何编码方式读取文件: {file_path}")
                    return None
            elif file_ext == ".pdf":
                return self._load_pdf_content(file_path)
            elif file_ext in {".docx", ".doc"}:
                return self._load_word_content(file_path)
            else:
                logging.warning(f"不支持的文件格式: {file_ext}")
                return None
        except Exception as e:
            logging.error(f"加载文件 {file_path} 失败: {e}")
            return None

    def _load_pdf_content(self, file_path: str) -> Optional[str]:
        """加载PDF文件内容"""
        try:
            import PyPDF2

            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                # 清理PDF文本，去除多余的空白字符
                text = self._clean_pdf_text(text)
                return text
        except ImportError:
            logging.error("请安装PyPDF2: pip install PyPDF2")
            return None
        except Exception as e:
            logging.error(f"PDF加载失败: {e}")
            return None

    def _clean_pdf_text(self, text: str) -> str:
        """清理PDF文本，去除多余的空白字符和换行符"""
        import re

        # 去除多余的空白字符
        text = re.sub(r"\s+", " ", text)
        # 去除行首行尾空白
        text = text.strip()
        # 将多个换行符替换为单个换行符
        text = re.sub(r"\n+", "\n", text)

        return text

    def _load_word_content(self, file_path: str) -> Optional[str]:
        """加载Word文件内容"""
        try:
            from docx import Document

            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            logging.error("请安装python-docx: pip install python-docx")
            return None
        except Exception as e:
            logging.error(f"Word文档加载失败: {e}")
            return None

    def _chunk_text(self, text: str) -> List[str]:
        """将文本分块"""
        chunk_size = 1000
        chunk_overlap = 200

        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            if end < len(text):
                sentence_end = max(
                    text.rfind(".", start, end),
                    text.rfind("!", start, end),
                    text.rfind("?", start, end),
                    text.rfind("\n", start, end),
                )

                if sentence_end > start:
                    end = sentence_end + 1

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = max(start + 1, end - chunk_overlap)

        return chunks

    def _generate_content_hash(self, content: str, file_path: Optional[str] = None) -> str:
        """生成文档内容的哈希值，pdf/doc/docx用二进制hash，其它用文本hash"""
        import hashlib
        import os
        from pathlib import Path

        # 如果是PDF或Word等二进制文件，直接用文件二进制内容hash
        if file_path and Path(file_path).suffix.lower() in {".pdf", ".docx", ".doc"} and os.path.exists(file_path):
            with open(file_path, "rb") as f:
                file_bytes = f.read()
            return hashlib.md5(file_bytes).hexdigest()
        # 否则用文本内容hash
        cleaned_content = self._clean_pdf_text(content).lower()
        return hashlib.md5(cleaned_content.encode("utf-8")).hexdigest()

    def build_query_workflow(self) -> Workflow:
        """构建查询工作流"""
        retrieval_config = self.config.get("retrieval", {})
        prompts = self.config.get("prompts", {})

        # 创建工作流上下文
        context = WorkflowContext()

        # 创建顶点
        source_vertex = SourceVertex(id="QUERY_SOURCE", name="查询输入", task=lambda inputs, context: inputs)

        # 创建查询改写的LLM顶点
        query_rewrite_llm = LLMVertex(
            id="QUERY_REWRITE_LLM",
            name="查询改写LLM",
            model=self.llm_model,
            params={
                SYSTEM: "你是一个查询改写和问题扩展专家。你可以使用网络搜索工具来帮助理解用户的问题，发现相关的热点讨论和扩展问题。\n\n你的任务是：\n1. 分析用户的原始问题\n2. 使用网络搜索了解这个问题的最新讨论和相关方面\n3. 基于搜索结果，生成多个不同角度的查询\n4. 提出一些值得探讨的扩展问题\n\n你必须以JSON格式返回结果。",
                USER: [
                    "请分析以下问题，并使用网络搜索来帮助扩展和改写查询。\n\n原始查询：{{original_query}}\n\n请以下面的JSON格式返回结果：\n{\n  \"original_query\": \"原始查询文本\",\n  \"intent_queries\": [\n    \"意图查询1（理解用户真正想问什么）\",\n    \"意图查询2（从不同角度理解）\",\n    \"意图查询3（考虑潜在需求）\"\n  ],\n  \"search_queries\": [\n    \"搜索查询1（关键信息检索）\",\n    \"搜索查询2（扩展信息检索）\"\n  ],\n  \"extended_questions\": [\n    \"扩展问题1（基于网络搜索发现的相关话题）\",\n    \"扩展问题2（值得深入探讨的方面）\",\n    \"扩展问题3（最新的相关讨论）\"\n  ]\n}\n\n确保返回的是合法的JSON格式。"
                ],
            },
            variables=[
                {
                    "source_scope": "QUERY_SOURCE",
                    "source_var": "query",
                    "local_var": "original_query",
                }
            ],
            tools=[self.web_search_service] if self.web_search_service else [],  # 添加web search工具
        )

        # 创建查询结果加工顶点
        query_processor = FunctionVertex(
            id="QUERY_PROCESSOR",
            name="查询结果加工",
            task=self._process_query_results,
            params={},
            variables=[
                {
                    "source_scope": "QUERY_REWRITE_LLM",
                    "source_var": None,
                    "local_var": "llm_output",
                }
            ],
        )

        # 创建查询改写的VertexGroup
        query_rewrite_group = VertexGroup(
            id="QUERY_REWRITE_GROUP",
            name="查询改写组",
            subgraph_vertices=[query_rewrite_llm, query_processor],
            subgraph_edges=[],
            variables=[
                {
                    "source_scope": None,
                    "source_var": "query",
                    "local_var": "original_query",
                },
                {
                    "source_scope": "QUERY_PROCESSOR",
                    "source_var": "processed_queries",
                    "local_var": "processed_queries",
                }
            ],
        )

        # 添加子图边：查询改写LLM -> 查询结果加工
        from .edge import Edge, EdgeType
        rewrite_to_processor_edge = Edge(
            source_vertex=query_rewrite_llm,
            target_vertex=query_processor,
            edge_type=Edge.ALWAYS,
        )
        query_rewrite_group.add_subgraph_edge(rewrite_to_processor_edge)

        # 创建查询向量化顶点（处理三种query）
        query_embedding = EmbeddingVertex(
            id="QUERY_EMBEDDING",
            name="查询向量化",
            embedding_provider=self.embedding_provider,
            variables=[
                {
                    "source_scope": "QUERY_REWRITE_GROUP",
                    "source_var": "processed_queries",
                    "local_var": "docs",
                }
            ],
        )

        # 创建向量查询顶点
        vector_query = VectorQueryVertex(
            id="VECTOR_QUERY",
            name="向量查询",
            vector_engine=self.vector_engine,
            params={
                "top_k": retrieval_config.get("top_k", 3),
                "similarity_threshold": retrieval_config.get("similarity_threshold", 0.3)
            },
            variables=[
                {
                    "source_scope": "QUERY_EMBEDDING",
                    "source_var": "embeddings",
                    "local_var": "query",
                }
            ],
        )

        # 创建Web搜索集成顶点
        web_search_vertex = WebSearchIntegrationVertex(
            id="WEB_SEARCH",
            name="Web搜索集成",
            params={
                "search_count": retrieval_config.get("web_search_count", 3),
                "enable_summary": retrieval_config.get("enable_web_summary", True)
            },
            web_search_service=self.web_search_service,
            variables=[
                {
                    "source_scope": "QUERY_REWRITE_GROUP",
                    "source_var": "processed_queries",
                    "local_var": "search_queries",
                },
                {
                    "source_scope": "QUERY_SOURCE",
                    "source_var": "query",
                    "local_var": "original_query",
                }
            ],
        )

        # 使用增强的结果格式化器，能够处理Web搜索结果
        result_formatter = EnhancedResultFormatterVertex(
            id="RESULT_FORMATTER",
            name="结果格式化",
            params={},
            variables=[
                {
                    "source_scope": "VECTOR_QUERY",
                    "source_var": "results",
                    "local_var": "results",
                },
                {
                    "source_scope": "WEB_SEARCH",
                    "source_var": "web_search_results",
                    "local_var": "web_search_results",
                },
                {
                    "source_scope": "WEB_SEARCH",
                    "source_var": "web_search_summary",
                    "local_var": "web_search_summary",
                },
                {
                    "source_scope": "WEB_SEARCH",
                    "source_var": "search_performed",
                    "local_var": "search_performed",
                }
            ],
        )

        # 创建最终LLM生成顶点
        llm_vertex = LLMVertex(
            id="LLM_GENERATE",
            name="生成回答",
            model=self.llm_model,
            params={
                SYSTEM: prompts.get("system", "你是一个有用的AI助手。你可以使用网络搜索工具来获取最新信息。当你需要查找实时或最新的信息时，请使用网络搜索工具。"),
                USER: [
                    "基于以下上下文信息回答问题：\n\n上下文：{{formatted_results}}\n\n问题：{{user_question}}\n\n扩展问题：{{extended_questions}}\n\n请提供准确、有用的回答。如果需要查找最新信息，可以使用网络搜索工具。"
                ],
            },
            tools=[self.web_search_service] if self.web_search_service else [],
            variables=[
                {
                    "source_scope": "RESULT_FORMATTER",
                    "source_var": "formatted_results",
                    "local_var": "formatted_results",
                },
                {
                    "source_scope": "QUERY_SOURCE",
                    "source_var": "query",
                    "local_var": "user_question",
                },
                {
                    "source_scope": "QUERY_REWRITE_GROUP",
                    "source_var": "processed_queries",
                    "local_var": "extended_questions",
                },
            ],
        )

        sink_vertex = SinkVertex(
            id="ANSWER_SINK",
            name="输出答案",
            task=lambda inputs, context: logging.info(f"生成答案: {inputs.get('answer', '')}"),
            variables=[
                {
                    "source_scope": "LLM_GENERATE",
                    "source_var": None,
                    "local_var": "answer",
                }
            ],
        )

        # 创建工作流
        workflow = Workflow(context=context)
        workflow.add_vertex(source_vertex)
        workflow.add_vertex(query_rewrite_group)
        workflow.add_vertex(query_embedding)
        workflow.add_vertex(vector_query)
        workflow.add_vertex(web_search_vertex)
        workflow.add_vertex(result_formatter)
        workflow.add_vertex(llm_vertex)
        workflow.add_vertex(sink_vertex)

        # 构建工作流图 - 并行执行向量查询和Web搜索
        source_vertex | query_rewrite_group
        query_rewrite_group | query_embedding | vector_query
        query_rewrite_group | web_search_vertex
        # 向量查询和Web搜索的结果都汇聚到结果格式化器
        vector_query | result_formatter
        web_search_vertex | result_formatter
        result_formatter | llm_vertex | sink_vertex

        return workflow

    def _process_query_results(self, inputs: Dict[str, Any], context=None):
        """
        处理LLM查询改写的结果，提取三种类型的查询
        
        Args:
            inputs: 包含LLM输出的输入
            
        Returns:
            包含原始查询、意图查询、搜索查询和扩展问题的结果
        """
        llm_output = inputs.get("llm_output", "")
        original_input_query = inputs.get("original_query", "")
        
        # 如果没有LLM输出，使用原始查询
        if not llm_output:
            return {
                "original_query": original_input_query or "未识别查询",
                "intent_queries": [],
                "search_queries": [],
                "extended_questions": [],  # 添加空的扩展问题列表
                "processed_queries": [original_input_query] if original_input_query else ["未识别查询"]
            }
        
        try:
            # 清理markdown代码块标记
            import re
            json_str = llm_output
            # 移除```json和```标记
            json_match = re.search(r'```(?:json)?\s*(.*?)```', json_str, re.DOTALL)
            if json_match:
                json_str = json_match.group(1).strip()
            
            # 尝试解析JSON输出
            import json
            result = json.loads(json_str)
            
            # 提取查询，使用原始输入作为后备
            original_query = result.get("original_query", "") or original_input_query or "未识别查询"
            intent_queries = [q for q in result.get("intent_queries", []) if q and isinstance(q, str)]
            search_queries = [q for q in result.get("search_queries", []) if q and isinstance(q, str)]
            extended_questions = [q for q in result.get("extended_questions", []) if q and isinstance(q, str)]  # 提取扩展问题
            
            # 合并所有查询用于向量化，确保至少包含原始查询
            all_queries = [q for q in ([original_query] + intent_queries + search_queries) if q and isinstance(q, str)]
            if not all_queries:
                all_queries = [original_input_query] if original_input_query else ["未识别查询"]
            
            return {
                "original_query": original_query,
                "intent_queries": intent_queries,
                "search_queries": search_queries,
                "extended_questions": extended_questions,  # 添加扩展问题到输出
                "processed_queries": all_queries
            }
            
        except (json.JSONDecodeError, AttributeError) as e:
            logging.error(f"解析LLM输出的JSON失败: {e}")
            logging.error(f"原始输出: {llm_output}")
            # 如果JSON解析失败，使用原始查询
            fallback_query = original_input_query or "未识别查询"
            return {
                "original_query": fallback_query,
                "intent_queries": [],
                "search_queries": [],
                "extended_questions": [],  # 添加空的扩展问题列表
                "processed_queries": [fallback_query]
            }

    def _merge_query_results(self, inputs: Dict[str, Any], context=None):
        """
        对每个查询向量进行检索并合并结果
        
        Args:
            inputs: 包含多个embedding的输入
            
        Returns:
            合并后的检索结果
        """
        embeddings = inputs.get("embeddings", [])
        if not embeddings:
            raise ValueError("No embeddings provided for search")
            
        # 获取检索参数
        top_k = self.config.get("retrieval", {}).get("top_k", 3)
        similarity_threshold = self.config.get("retrieval", {}).get("similarity_threshold", 0.3)
        
        all_results = []
        seen_docs = set()  # 用于去重
        
        # 对每个embedding进行检索
        for embedding in embeddings:
            try:
                results = self.vector_engine.search(
                    embedding, 
                    top_k=top_k
                )
                
                # 过滤低于阈值的结果并去重
                for result in results:
                    doc_id = result.get("id")
                    score = result.get("score", 0)
                    if score >= similarity_threshold and doc_id not in seen_docs:
                        seen_docs.add(doc_id)
                        all_results.append(result)
                        
            except Exception as e:
                logging.error(f"单个查询检索失败: {e}")
                continue
        
        # 按相似度分数排序
        all_results.sort(key=lambda x: x.get("score", 0), reverse=True)
        
        # 只返回top_k个结果
        final_results = all_results[:top_k]
        
        if not final_results:
            logging.warning("所有查询都未返回结果")
            return {"results": []}
            
        return {"results": final_results}


class UnifiedRAGSystem:
    """统一配置的RAG系统"""

    def __init__(self):
        """
        初始化统一配置的RAG系统

        """
        # 检查RAG依赖是否已安装
        self._check_rag_dependencies()

        self.builder = UnifiedRAGWorkflowBuilder()

        self.indexing_workflow = None
        self.query_workflow = None

        # 性能优化: 添加缓存和复用机制
        self._query_workflow_instance = None
        self._embedding_cache = {}  # 查询embedding缓存
        self._is_initialized = False
        self._indexing_only_mode = False  # 仅索引模式标志

    def _check_rag_dependencies(self):
        """检查RAG依赖是否已安装"""
        missing_deps = []

        try:
            import sentence_transformers
        except ImportError:
            missing_deps.append("sentence-transformers")

        try:
            import faiss
        except ImportError:
            missing_deps.append("faiss-cpu")

        try:
            import numpy
        except ImportError:
            missing_deps.append("numpy")

        if missing_deps:
            error_msg = f"RAG功能需要以下依赖包，但未安装: {', '.join(missing_deps)}\n"
            error_msg += "请使用以下命令安装RAG依赖:\n"
            error_msg += "  uv pip install vertex[rag]\n"
            error_msg += "或者手动安装:\n"
            error_msg += "  uv pip install sentence-transformers faiss-cpu numpy"
            raise ImportError(error_msg)

    def _lazy_initialize(self, indexing_only: bool = False):
        """
        延迟初始化，只在需要时构建工作流

        Args:
            indexing_only: 是否仅需要索引功能（不初始化LLM）
        """
        if not self._is_initialized:
            self._indexing_only_mode = indexing_only
            if indexing_only:
                # 仅构建索引工作流，不初始化LLM
                self.indexing_workflow = self.builder.build_smart_indexing_workflow()
                logging.info("仅索引模式：跳过LLM初始化")
            else:
                # 完整初始化
                self.build_workflows()
            self._is_initialized = True

    def build_workflows(self):
        """构建索引和查询工作流"""
        self.indexing_workflow = self.builder.build_smart_indexing_workflow()

        # 只有在非仅索引模式下才构建查询工作流
        if not self._indexing_only_mode:
            try:
                self.query_workflow = self.builder.build_query_workflow()
                # 性能优化: 预构建可复用的查询工作流实例
                self._query_workflow_instance = self.builder.build_query_workflow()
                logging.info("统一配置RAG工作流构建完成")
            except Exception as e:
                logging.warning(f"查询工作流构建失败，可能是LLM配置问题: {e}")
                logging.info("系统将以仅索引模式运行")
                self._indexing_only_mode = True

    def get_vector_db_stats(self):
        """获取向量数据库统计信息"""
        if hasattr(self.builder, "vector_engine") and self.builder.vector_engine:
            if hasattr(self.builder.vector_engine, "get_stats"):
                return self.builder.vector_engine.get_stats()
            else:
                return {"message": "向量引擎不支持统计信息"}
        return {"message": "向量引擎未初始化"}

    def _get_cached_embedding(self, text: str):
        """获取缓存的查询embedding"""
        import hashlib

        text_hash = hashlib.md5(text.encode("utf-8")).hexdigest()
        return self._embedding_cache.get(text_hash)

    def _cache_embedding(self, text: str, embedding):
        """缓存查询embedding"""
        import hashlib

        text_hash = hashlib.md5(text.encode("utf-8")).hexdigest()
        self._embedding_cache[text_hash] = embedding

        # 限制缓存大小，避免内存溢出
        if len(self._embedding_cache) > 100:
            # 删除最旧的缓存项
            oldest_key = next(iter(self._embedding_cache))
            del self._embedding_cache[oldest_key]

    def index_documents(self, documents: List[str], force_reindex: bool = False, update_existing: bool = True):
        """
        索引文档

        Args:
            documents: 文档列表，可以是文件路径或文档内容
            force_reindex: 是否强制重新索引，默认False
            update_existing: 是否更新已存在的文档，默认True
        """
        # 仅索引模式初始化，不加载LLM
        self._lazy_initialize(indexing_only=True)

        # 检查向量数据库状态
        stats = self.get_vector_db_stats()
        existing_docs = stats.get("total_documents", 0)

        if isinstance(existing_docs, (int, float)) and existing_docs > 0 and not force_reindex and not update_existing:
            print(
                f"向量数据库中已有 {existing_docs} 个文档，跳过索引。使用 force_reindex=True 强制重新索引，或 update_existing=True 更新现有文档。"
            )
            return

        if isinstance(existing_docs, (int, float)) and existing_docs > 0 and update_existing:
            print(f"向量数据库中已有 {existing_docs} 个文档，将检测并更新变化的文档...")
        else:
            logging.info(f"开始索引 {len(documents)} 个文档")

        # 每次执行都创建新的工作流实例，避免重复运行错误
        workflow_instance = self.builder.build_smart_indexing_workflow()
        workflow_instance.execute_workflow(source_inputs={"docs": documents})

        # 显示索引后的统计信息
        new_stats = self.get_vector_db_stats()
        new_doc_count = new_stats.get("total_documents", 0)
        logging.info(f"文档索引完成，向量数据库现在包含 {new_doc_count} 个文档")

        if update_existing and isinstance(existing_docs, (int, float)) and existing_docs > 0:
            if isinstance(new_doc_count, (int, float)):
                added_docs = new_doc_count - existing_docs
                if added_docs > 0:
                    print(f"更新完成：新增了 {added_docs} 个文档")
                else:
                    print("更新完成：没有新增文档（可能都是重复内容）")
        
        # 重要：索引完成后，重置初始化状态，确保后续查询可以正常使用LLM
        # 但保持已初始化的组件（embedding_provider, vector_engine等）
        if self._indexing_only_mode:
            logging.info("索引完成，重置为完整模式以支持LLM查询")
            self._indexing_only_mode = False
            # 重新构建查询工作流
            try:
                self.query_workflow = self.builder.build_query_workflow()
                self._query_workflow_instance = self.builder.build_query_workflow()
                logging.info("查询工作流重新构建完成")
            except Exception as e:
                logging.warning(f"查询工作流重新构建失败: {e}")
                self._indexing_only_mode = True

    def query(self, question: str, use_cache: bool = True) -> str:
        """
        查询问题（性能优化版本）

        Args:
            question: 问题
            use_cache: 是否使用embedding缓存，默认True

        Returns:
            生成的答案
        """
        # 查询需要完整初始化（包括LLM）
        self._lazy_initialize(indexing_only=False)

        # 检查是否在仅索引模式
        if self._indexing_only_mode:
            logging.warning("系统在仅索引模式下运行，无法执行LLM查询，回退到快速查询模式")
            return self.query_fast(question)

        logging.info(f"查询问题: {question}")

        # 性能优化: 复用工作流实例，避免重复构建
        if self._query_workflow_instance is None:
            try:
                self._query_workflow_instance = self.builder.build_query_workflow()
            except Exception as e:
                logging.error(f"查询工作流构建失败: {e}")
                return self.query_fast(question)

        try:
            # 清理之前的输出状态
            self._reset_workflow_state(self._query_workflow_instance)

            retrieval_config = self.builder.config.get("retrieval", {}) if isinstance(self.builder.config, dict) else {}

            self._query_workflow_instance.execute_workflow(
                source_inputs={
                    "query": question,
                    "top_k": retrieval_config.get("top_k", 3) if isinstance(retrieval_config, dict) else 3,
                    "similarity_threshold": (
                        retrieval_config.get("similarity_threshold", 0.3) if isinstance(retrieval_config, dict) else 0.3
                    ),
                }
            )

            # 获取LLM的输出
            llm_vertex = self._query_workflow_instance.get_vertice_by_id("LLM_GENERATE")
            if llm_vertex and hasattr(llm_vertex, "output"):
                return llm_vertex.output if llm_vertex.output else "无法生成答案"
            else:
                return "无法生成答案"

        except Exception as e:
            logging.error(f"查询执行失败: {e}")
            # 如果复用的工作流失败，尝试创建新的
            try:
                workflow_instance = self.builder.build_query_workflow()
                retrieval_config = (
                    self.builder.config.get("retrieval", {}) if isinstance(self.builder.config, dict) else {}
                )

                workflow_instance.execute_workflow(
                    source_inputs={
                        "query": question,
                        "top_k": retrieval_config.get("top_k", 3) if isinstance(retrieval_config, dict) else 3,
                        "similarity_threshold": (
                            retrieval_config.get("similarity_threshold", 0.3)
                            if isinstance(retrieval_config, dict)
                            else 0.3
                        ),
                    }
                )

                llm_vertex = workflow_instance.get_vertice_by_id("LLM_GENERATE")
                if llm_vertex and hasattr(llm_vertex, "output"):
                    return llm_vertex.output if llm_vertex.output else "无法生成答案"
                else:
                    return "无法生成答案"
            except Exception as e2:
                logging.error(f"查询彻底失败，回退到快速查询: {e2}")
                return self.query_fast(question)

    def _reset_workflow_state(self, workflow_instance):
        """重置工作流状态，准备下次执行"""
        if workflow_instance and hasattr(workflow_instance, "vertices"):
            for vertex in workflow_instance.vertices.values():
                if hasattr(vertex, "output"):
                    vertex.output = None
                if hasattr(vertex, "_executed"):
                    vertex._executed = False

    def query_fast(self, question: str) -> str:
        """
        快速查询模式（最小化LLM调用）

        Args:
            question: 问题

        Returns:
            格式化的检索结果或LLM答案
        """
        # 快速查询只需要索引组件，不需要LLM
        self._lazy_initialize(indexing_only=True)

        # 直接进行向量检索，跳过LLM生成
        try:
            # 获取查询embedding
            if self.builder.embedding_provider:
                query_embedding = self.builder.embedding_provider.embedding(question)
                if query_embedding:
                    # 直接搜索向量数据库
                    retrieval_config = (
                        self.builder.config.get("retrieval", {}) if isinstance(self.builder.config, dict) else {}
                    )
                    top_k = retrieval_config.get("top_k", 3) if isinstance(retrieval_config, dict) else 3

                    results = self.builder.vector_engine.search(query_embedding, top_k=top_k, include_vector=False)

                    # 格式化结果
                    if results:
                        formatted_results = []
                        for i, result in enumerate(results, 1):
                            content = result.get("content", "")
                            score = result.get("score", 0)
                            doc_id = result.get("id", f"文档{i}")
                            formatted_result = f"文档{i} (ID: {doc_id}, 相似度: {score:.3f}):\n{content}\n"
                            formatted_results.append(formatted_result)
                        return "\n".join(formatted_results)
                    else:
                        return "未找到相关文档"
                else:
                    return "查询embedding生成失败"
            else:
                return "embedding提供者未初始化"
        except Exception as e:
            logging.error(f"快速查询失败: {e}")
            return f"查询失败: {e}"

    def start_interactive_mode(self, fast_mode: bool = False):
        """
        启动交互式查询模式（性能优化版本）

        Args:
            fast_mode: 是否使用快速模式（跳过LLM生成）
        """
        # 根据模式选择初始化方式
        self._lazy_initialize(indexing_only=fast_mode)

        print("\n=== 交互式查询模式 ===")
        if fast_mode:
            print("(快速模式 - 仅显示检索结果，不调用LLM)")
        else:
            if self._indexing_only_mode:
                print("(检测到LLM配置问题，自动切换到快速模式)")
                fast_mode = True
        print("输入 'quit', 'exit' 或 '退出' 来结束")
        print("输入 'stats' 查看数据库统计")
        print("输入 'cache' 查看缓存状态")
        print("=" * 50)

        query_count = 0
        start_time = time.time()

        while True:
            try:
                question = input("\n请输入您的问题: ").strip()

                if question.lower() in ["quit", "exit", "退出"]:
                    break
                elif question.lower() == "stats":
                    stats = self.get_vector_db_stats()
                    print(f"\n数据库统计: {stats}")
                    continue
                elif question.lower() == "cache":
                    print(f"\n缓存状态: {len(self._embedding_cache)} 个查询已缓存")
                    continue

                if not question:
                    continue

                query_start = time.time()
                print("\n正在生成答案...")

                if fast_mode or self._indexing_only_mode:
                    answer = self.query_fast(question)
                else:
                    answer = self.query(question)

                query_time = time.time() - query_start
                query_count += 1

                print(f"\n答案: {answer}")
                print(f"查询耗时: {query_time:.2f}秒")
                print("-" * 50)

            except KeyboardInterrupt:
                print("\n\n已取消查询")
                break
            except EOFError:
                break

        total_time = time.time() - start_time
        if query_count > 0:
            avg_time = total_time / query_count
            print(f"\n会话统计: {query_count} 次查询，平均耗时 {avg_time:.2f}秒")

    def show_workflows(self):
        """显示工作流图"""
        if self.indexing_workflow:
            print("=== 索引工作流 ===")
            self.indexing_workflow.show_graph()

        if self.query_workflow:
            print("\n=== 查询工作流 ===")
            self.query_workflow.show_graph()
        elif self._indexing_only_mode:
            print("\n=== 查询工作流 ===")
            print("(仅索引模式 - 查询工作流未构建)")


def main():
    """示例用法"""
    # 创建示例文档
    sample_docs = [
        "人工智能是计算机科学的一个分支，致力于创建能够执行通常需要人类智能的任务的系统。",
        "机器学习是人工智能的一个子集，它使计算机能够在没有明确编程的情况下学习和改进。",
        "深度学习是机器学习的一个分支，使用神经网络来模拟人脑的学习过程。",
        "自然语言处理是人工智能的一个领域，专注于计算机理解和生成人类语言的能力。",
        "计算机视觉是人工智能的一个分支，使计算机能够从图像和视频中理解和提取信息。",
    ]

    # 创建统一配置的RAG系统
    rag_system = UnifiedRAGSystem()

    # 构建工作流
    rag_system.build_workflows()

    # 显示工作流图
    rag_system.show_workflows()

    # 索引文档
    rag_system.index_documents(sample_docs)

    # 查询示例
    questions = ["什么是人工智能？", "机器学习和深度学习有什么区别？", "自然语言处理有什么应用？"]

    for question in questions:
        print(f"\n问题: {question}")
        answer = rag_system.query(question)
        print(f"答案: {answer}")


if __name__ == "__main__":
    main()
